from argparse import ArgumentParser
from dataclasses import dataclass
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from copy import deepcopy

START_TITLE_LIST = [
    "Microsoft Teams for Win32 no GB checkpoint",
    "Microsoft Teams for Mac no GB checkpoint",
    "Microsoft Teams for Online no GB checkpoint",
    "Microsoft Teams for Android no GB checkpoint",
    "Microsoft Teams for iOS no GB checkpoint",
]


@dataclass
class DocumentPart2:
    title: Paragraph
    overview_en: list[Paragraph]
    overview_zh: list[Paragraph]


@dataclass
class DocumentPart:
    title: Paragraph
    part_list: list[DocumentPart2]


def _get_xml_part_from_paragraph(paragraph: Paragraph) -> list[object]:
    return paragraph._element.xpath(
        ".//w:r | .//w:commentRangeStart | .//w:commentRangeEnd"
    )


def _get_xml_text(run_xml) -> str:
    try:
        return "".join([run.text for run in run_xml.xpath(".//w:t")])
    except Exception:
        return None


def txt_lines_to_part2(lines: list[Paragraph]) -> DocumentPart2:
    title = lines[0]
    # Find zh title under startwith 中文翻译 and underlines have 2
    for i in range(1, len(lines)):
        if lines[i].text.strip().startswith("中文翻译"):
            if len(lines) - i >= 3:
                title = lines[i + 1]
            overview_zh = lines[i:]
            overview_en = lines[:i]
            break
    return DocumentPart2(title=title, overview_en=overview_en, overview_zh=overview_zh)


def _remove_number_from_paragraph(paragraph: Paragraph) -> Paragraph:
    m = re.match(r"^\d+\)", paragraph.text.strip())
    number_text = m.group(0) if m else ""
    if m:
        for run in paragraph.runs:
            while number_text and run.text:
                run.text = run.text.lstrip()
                char = number_text[0]
                if run.text.startswith(char):
                    run.text = run.text[1:]
                    number_text = number_text[1:]
    return paragraph


chinese_pattern = re.compile(r"[\u4e00-\u9fff]")


def txt_parse_document_part(paragraphs: list[Paragraph]) -> list[DocumentPart]:
    lines = [line for line in paragraphs if line.text.strip()]
    part1_list = []

    tmp_list = []
    for i in range(len(lines)):
        line = lines[i]
        create_new_part = False
        create_new_part2 = False
        # Check line is start with number)
        # And get the number
        if line.text.strip() in START_TITLE_LIST:
            create_new_part = True
            create_new_part2 = True
        elif line._element.xpath(".//w:ilvl"):
            if tmp_list and not (
                chinese_pattern.search(lines[i - 1].text)
                and not chinese_pattern.search(line.text)
            ):
                # Look like is in the same part
                pass
            else:
                create_new_part2 = True
            # Remove d) from the line
        if i >= len(lines) - 1:
            tmp_list.append(line)
            create_new_part2 = True
        if create_new_part2 and tmp_list:
            part1_list[-1].part_list.append(txt_lines_to_part2(tmp_list))
            tmp_list = []
        if create_new_part:
            part1_list.append(DocumentPart(line, part_list=[]))
        else:
            tmp_list.append(line)
    return part1_list


PARAGRAPH_XML_MAP = {
    "title1": None,
    "title2": None,
    "overview_title": None,
    "overview_en": None,
    "overview_zh": None,
    "overview_end1": None,
    "overview_end2": None,
}


def doc_init_xml_map(paragraphs: list[Paragraph]):
    for p in paragraphs:
        if p.text == "title1":
            PARAGRAPH_XML_MAP["title1"] = p
        elif p.text == "title2":
            PARAGRAPH_XML_MAP["title2"] = p
        elif p.text == "概述":
            PARAGRAPH_XML_MAP["overview_title"] = p
        elif p.text == "eng_body":
            PARAGRAPH_XML_MAP["overview_en"] = p
        elif p.text == "zh_body":
            PARAGRAPH_XML_MAP["overview_zh"] = p
        elif p.text == "测试场景":
            PARAGRAPH_XML_MAP["overview_end1"] = p
        elif p.text == "此新功能没有GB18030相关的测试场景。":
            PARAGRAPH_XML_MAP["overview_end2"] = p


def _clear_content(element):
    """Remove all child elements"""
    for e in element.xpath("./*"):
        element.remove(e)


def doc_add_paragraph(
    doc: Document,
    para_type: str,
    part_list: Paragraph | list[Paragraph] | None = None,
):
    if not isinstance(part_list, list):
        part_list = [part_list]
    for part in part_list:
        paragraph = doc.add_paragraph()
        _clear_content(paragraph._element)
        items = PARAGRAPH_XML_MAP[para_type]._element.xpath("./*")
        items = deepcopy(items)
        for item in items:
            paragraph._element.append(item)
        run = paragraph.runs[0]
        if part is not None:
            run_xml_temp = deepcopy(run._element)
            # Remove all run elements
            paragraph._element.remove(run._element)
            run_xml_list = _get_xml_part_from_paragraph(part)
            run_xml_temp = deepcopy(run_xml_temp)
            for run_xml in run_xml_list:
                text = _get_xml_text(run_xml)
                if text:
                    paragraph._element.append(deepcopy(run_xml_temp))
                    run = paragraph.runs[-1]
                    run.text = text
                else:
                    paragraph._element.append(run_xml)


def main():
    parser = ArgumentParser(description="Process some files.")
    parser.add_argument("--temp", required=True, help="Path to the template file")
    parser.add_argument("--doc", required=True, help="Path to the output file")
    args = parser.parse_args()
    doc_path = args.temp
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    content_start_index = None
    for i in range(len(paragraphs)):
        if paragraphs[i].text == "===":
            content_start_index = i
            break
    doc_init_xml_map(paragraphs[:content_start_index])
    part_list = txt_parse_document_part(paragraphs[content_start_index + 1 :])
    doc.add_paragraph("===")
    for part in part_list:
        doc_add_paragraph(doc, "title1", part.title)
        for part2 in part.part_list:
            doc_add_paragraph(doc, "title2", part2.title)
            doc_add_paragraph(doc, "overview_title")
            doc_add_paragraph(doc, "overview_en", part2.overview_en)
            doc_add_paragraph(doc, "overview_zh", part2.overview_zh)
            doc_add_paragraph(doc, "overview_end1")
            doc_add_paragraph(doc, "overview_end2")
    doc.save(args.doc)


if __name__ == "__main__":
    main()
